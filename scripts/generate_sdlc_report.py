"""
Generate SDLC-Report.docx for EPMCDMETST-40786 — Borrow a Book (FR-3.1)
Run: backend\.venv\Scripts\python scripts\generate_sdlc_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT  = os.path.join(ROOT, "SDLC-Report.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        tc   = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  header_color)
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for ci, val in enumerate(row_data):
            row_cells[ci].text = str(val)
            row_cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if ri % 2 == 0:
            shade_row(row_cells[0].rows[0] if False else table.rows[ri + 1], "DCE6F1")

    # Column widths
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()
    return table


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p


def body(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold  = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def badge(doc, label, value, ok=True):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(f"{value}  {'✅ PASS' if ok else '❌ FAIL'}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x37, 0x86, 0x10) if ok else RGBColor(0xC0, 0x00, 0x00)
    return p


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ---------------------------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------------------------

doc.add_paragraph("\n\n")
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Online Library Management System")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("SDLC Execution Report")
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph("\n")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Feature:  FR-3.1 — Borrow a Book\n").bold = True
meta.add_run("Jira Ticket:  EPMCDMETST-40786\n").bold = True
meta.add_run("Date:  June 4, 2026\n").bold = True
meta.add_run("Branch:  EPMCDMETST-40786\n").bold = True
meta.add_run("PR:  #3 — https://github.com/JanpreetSingh/online-library-management/pull/3").bold = True

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 1 — APPLICATION OVERVIEW
# ---------------------------------------------------------------------------

h1(doc, "1. Application Overview")
body(doc,
    "The Online Library Management System is a full-stack web application for managing a "
    "library's books, users, and borrowing. It is built with React 18 + TypeScript on the "
    "frontend, FastAPI + SQLAlchemy + PostgreSQL on the backend, and deployed via Docker Compose.")

h2(doc, "1.1 Technology Stack")
add_table(doc,
    ["Layer", "Technology"],
    [
        ["Frontend",     "React 18 + TypeScript + Vite + Tailwind CSS"],
        ["Backend",      "Python FastAPI 0.115 + SQLAlchemy 2.0 + Pydantic v2"],
        ["Database",     "PostgreSQL 16 (Alpine)"],
        ["Auth",         "JWT (python-jose) + bcrypt (passlib)"],
        ["Containers",   "Docker Compose — db / backend / frontend services"],
        ["E2E Tests",    "Playwright"],
        ["Unit Tests",   "pytest + httpx (TestClient)"],
    ],
    col_widths=[1.5, 4.5]
)

h2(doc, "1.2 Roles & Access")
add_table(doc,
    ["Role", "Permissions"],
    [
        ["admin",     "Full access: manage users, books, transactions"],
        ["librarian", "Add/edit books, manage borrowing"],
        ["member",    "Browse books, borrow, view own transactions"],
        ["guest",     "Read-only browse — no borrowing"],
    ],
    col_widths=[1.2, 4.8]
)

h2(doc, "1.3 API Endpoints")
add_table(doc,
    ["Method", "Path", "Description", "Roles"],
    [
        ["POST", "/api/auth/login",          "Email + password → JWT",          "Public"],
        ["POST", "/api/auth/guest-login",    "Guest access → short-lived JWT",  "Public"],
        ["POST", "/api/auth/register",       "Create user",                     "Admin"],
        ["GET",  "/api/users/me",            "Current user profile",            "Authenticated"],
        ["PUT",  "/api/users/me",            "Update profile",                  "Authenticated"],
        ["GET",  "/api/books",               "List all books",                  "All"],
        ["GET",  "/api/books/{id}",          "Get single book",                 "All"],
        ["POST", "/api/books",               "Add a book",                      "Admin, Librarian"],
        ["PUT",  "/api/books/{id}",          "Edit a book",                     "Admin, Librarian"],
        ["POST", "/api/borrow/{book_id}",    "Borrow a book",                   "Member only"],
    ],
    col_widths=[0.7, 2.2, 2.5, 1.6]
)

h2(doc, "1.4 Database Schema")
add_table(doc,
    ["Table", "Key Columns", "Notes"],
    [
        ["users",              "id (UUID), name, email, role, is_active",       "role ENUM: admin/librarian/member/guest"],
        ["books",              "id (UUID), title, author, isbn, available_copies, total_copies", "available_copies ≥ 0 enforced at DB level"],
        ["borrow_transactions","id (UUID), user_id, book_id, borrowed_at, due_date, status",    "Indexes on user_id, book_id, status"],
    ],
    col_widths=[1.5, 3.2, 2.3]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 2 — SDLC ORCHESTRATOR FLOW
# ---------------------------------------------------------------------------

h1(doc, "2. SDLC Orchestrator Flow")
body(doc,
    "The project uses three orchestrator agents, each responsible for a distinct phase of the "
    "software delivery lifecycle. They are invoked in sequence by the developer.")

h2(doc, "2.1 Three-Orchestrator Pipeline")
add_table(doc,
    ["Orchestrator", "Agent Name", "Responsibility", "Handoff"],
    [
        ["BA Orchestrator",         "orchestrator-ba",         "Fetch requirements → Gap analysis → Create Jira stories", "→ orchestrator-sdlc <ticket>"],
        ["SDLC Orchestrator",       "orchestrator-sdlc",       "Refine Req → Design → Review → Plan → Code → Review → Test → PR", "→ orchestrator-deployment <ticket>"],
        ["Deployment Orchestrator", "orchestrator-deployment", "Deploy locally → Update Confluence docs", "Session end"],
    ],
    col_widths=[1.6, 1.8, 2.8, 1.8]
)

h2(doc, "2.2 orchestrator-sdlc — 8-Step Workflow")
body(doc, "Each step delegates to a specialist subagent and ends with a mandatory human approval gate.")
add_table(doc,
    ["Step", "Name", "Subagent", "Gate / Output"],
    [
        ["1", "Refine Requirements",   "requirements-assistant",  "REQUIREMENTS.md — human approval"],
        ["2", "Architecture Design",   "architecture-design",     "ARCHITECTURE.md — human approval"],
        ["3", "Design Review",         "design-review",           "design-review.md — ✅/❌ gate"],
        ["4", "Implementation Planning","implementation-planning", "implementation-plan.md — human approval"],
        ["5", "Implementation",        "implementation",          "git diff — human approval"],
        ["6", "Code Review",           "code-review-assistant",   "code-review.md — human approval"],
        ["7", "Verify",                "verify-test",             "verify-test-result.md — human approval"],
        ["8", "PR Creation",           "pr-creator",              "PR URL + CHANGELOG.md"],
    ],
    col_widths=[0.4, 1.8, 1.9, 3.0]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 3 — AGENTS IN THE ORCHESTRATOR
# ---------------------------------------------------------------------------

h1(doc, "3. Agents in the Orchestrator")
body(doc, "The orchestrator-sdlc uses 8 specialist subagents. Each has a focused role, minimal tool set, and produces a specific output artifact.")

add_table(doc,
    ["Agent", "Role", "Tools", "Output Artifact"],
    [
        ["requirements-assistant",  "Refine user story into full requirements doc",           "read, search, edit",                       "REQUIREMENTS.md"],
        ["architecture-design",     "Produce architecture overview + API contract",            "read, search, edit, confluence/*",         "ARCHITECTURE.md"],
        ["design-review",           "Review ARCHITECTURE.md for risks and gaps",              "read, search, edit",                       "design-review.md"],
        ["implementation-planning", "Break design into prioritised task list",                "read, search, edit",                       "implementation-plan.md"],
        ["implementation",          "Implement all tasks from implementation-plan.md",        "read, search, edit, execute",              "Code changes + git diff"],
        ["code-review-assistant",   "Review uncommitted diff for security/correctness",       "read, search, execute",                    "code-review.md"],
        ["verify-test",             "Run unit tests + write + execute Playwright E2E tests",  "read, edit, execute, playwright/*",        "verify-test-result.md"],
        ["pr-creator",              "Create PR with full description + changelog entry",      "read, search, execute, github/*",          "PR URL + CHANGELOG.md"],
    ],
    col_widths=[1.8, 2.1, 1.8, 1.8]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 4 — STAGE OUTPUTS
# ---------------------------------------------------------------------------

h1(doc, "4. Stage Outputs — EPMCDMETST-40786")
body(doc, "The following subsections document the actual output produced at each step of the orchestrator execution for FR-3.1: Borrow a Book.")

# --- STEP 1 ---
h2(doc, "Step 1 — Refine Requirements")
body(doc, "Agent: requirements-assistant     Output: REQUIREMENTS.md     Gate: ✅ Approved", italic=True)
add_table(doc,
    ["Item", "Detail"],
    [
        ["User Story",       "As a library member, I want to borrow an available book from the catalog, so that I can take it out on loan and the system accurately tracks availability."],
        ["Jira",             "EPMCDMETST-40786"],
        ["Functional Reqs",  "8 (authentication, availability check, borrow limit, role guard, transaction creation, copy decrement, atomicity, concurrency)"],
        ["Non-Functional",   "6 (security JWT, server-side auth, DB-level atomicity, no negative copies, error clarity, idempotency)"],
        ["Acceptance Criteria", "6 (happy path, zero copies, borrow limit, decrement invariant, role guard, concurrent request)"],
        ["Key Constraint",   "Max 5 active borrows per member; 409 Conflict on capacity/limit exceeded; 403 on role violation"],
        ["Out of Scope",     "Return flow, due dates, overdue handling, notifications, fines, reservations"],
    ],
    col_widths=[1.8, 5.2]
)

# --- STEP 2 ---
h2(doc, "Step 2 — Architecture Design")
body(doc, "Agent: architecture-design     Output: ARCHITECTURE.md, docs/FR-3.1-IMPLEMENTATION.md     Gate: ✅ Approved", italic=True)
add_table(doc,
    ["Component", "Design Decision"],
    [
        ["New endpoint",        "POST /api/borrow/{book_id}  —  member role only"],
        ["New model",           "BorrowTransaction: id, user_id, book_id, borrowed_at, due_date, returned_at, status"],
        ["Concurrency",         "SELECT FOR UPDATE on book row inside a DB transaction to prevent over-booking"],
        ["Constraint",          "CheckConstraint('available_copies >= 0') added to books table"],
        ["Auth pattern",        "HTTPBearer(auto_error=False) → 401 for missing/invalid JWT; 403 for wrong role"],
        ["Schema",              "BorrowBookResponse — Pydantic v2 with ConfigDict(from_attributes=True)"],
        ["Frontend",            "Borrow button in BooksPage.tsx visible to members only; inline service borrowService.ts"],
        ["DB indexes",          "borrow_transactions(user_id), (book_id), (status)"],
    ],
    col_widths=[1.8, 5.2]
)

# --- STEP 3 ---
h2(doc, "Step 3 — Design Review")
body(doc, "Agent: design-review     Output: design-review.md     Gate: ✅ Approved — no blocking issues", italic=True)
add_table(doc,
    ["Area", "Finding", "Severity"],
    [
        ["Concurrency",    "SELECT FOR UPDATE covers last-copy race condition correctly",   "✅ OK"],
        ["Auth",           "HTTPBearer auto_error=False ensures proper 401 before 403",    "✅ OK"],
        ["DB Constraint",  "available_copies ≥ 0 enforced at DB level as safety net",      "✅ OK"],
        ["Borrow limit",   "Checked inside the transaction — consistent with isolation",   "✅ OK"],
        ["Error codes",    "409 for capacity/limit, 403 for role, 401 for auth — clear",   "✅ OK"],
        ["Scope",          "Return flow correctly deferred to FR-3.2",                     "✅ OK"],
    ],
    col_widths=[1.6, 4.2, 1.2]
)

# --- STEP 4 ---
h2(doc, "Step 4 — Implementation Planning")
body(doc, "Agent: implementation-planning     Output: implementation-plan.md     Gate: ✅ Approved — 10 tasks", italic=True)
add_table(doc,
    ["Task", "Description", "Priority"],
    [
        ["TASK-1",  "Create BorrowTransaction SQLAlchemy model",                        "P1"],
        ["TASK-2",  "Register model in models/__init__.py + update seed.py",            "P1"],
        ["TASK-3",  "Create BorrowBookResponse Pydantic v2 schema",                    "P1"],
        ["TASK-4",  "Implement POST /api/borrow/{book_id} router with role guard",     "P1"],
        ["TASK-5",  "Add SELECT FOR UPDATE concurrency lock + atomic decrement",        "P1"],
        ["TASK-6",  "Register borrow router in main.py",                               "P1"],
        ["TASK-7",  "Update HTTPBearer to auto_error=False in dependencies.py",        "P2"],
        ["TASK-8",  "Add CheckConstraint on books.available_copies",                   "P2"],
        ["TASK-9",  "Create borrowService.ts + BorrowTransaction TypeScript type",     "P2"],
        ["TASK-10", "Add Borrow button to BooksPage.tsx (member only, UUID path param)","P2"],
    ],
    col_widths=[0.8, 4.5, 0.7]
)

# --- STEP 5 ---
h2(doc, "Step 5 — Implementation")
body(doc, "Agent: implementation     All 10 tasks completed ✓     Gate: ✅ Approved", italic=True)
add_table(doc,
    ["File", "Change Type", "Description"],
    [
        ["backend/app/models/borrow_transaction.py",  "NEW",      "BorrowTransaction ORM model with all columns + indexes"],
        ["backend/app/schemas/borrow_transaction.py", "NEW",      "BorrowBookResponse Pydantic v2 schema"],
        ["backend/app/routers/borrow.py",             "NEW",      "POST /api/borrow/{book_id} with role guard + SELECT FOR UPDATE"],
        ["backend/app/models/__init__.py",            "MODIFIED", "Import BorrowTransaction for table creation"],
        ["backend/app/main.py",                       "MODIFIED", "Register borrow router"],
        ["backend/app/auth/dependencies.py",          "MODIFIED", "HTTPBearer(auto_error=False) for proper 401/403 split"],
        ["backend/app/models/book.py",                "MODIFIED", "CheckConstraint('available_copies >= 0')"],
        ["backend/seed.py",                           "MODIFIED", "Import borrow model for table auto-creation"],
        ["frontend/src/types/borrowTransaction.ts",   "NEW",      "BorrowTransaction TypeScript type"],
        ["frontend/src/services/borrowService.ts",    "NEW",      "borrowBook(bookId) Axios service"],
        ["frontend/src/pages/BooksPage.tsx",          "MODIFIED", "Borrow button in book cards, member-only, loading state"],
        ["tests/e2e/deployment-verification.spec.ts", "MODIFIED", "Added 422 + 403 to expected status set for borrow endpoint"],
    ],
    col_widths=[3.2, 1.0, 2.8]
)

# --- STEP 6 ---
h2(doc, "Step 6 — Code Review")
body(doc, "Agent: code-review-assistant     Output: code-review.md     Gate: ✅ Approved — no blocking issues", italic=True)
add_table(doc,
    ["Area", "Finding", "Verdict"],
    [
        ["Security — JWT",        "All protected routes require valid JWT; 401 correctly returned before business logic", "✅ Pass"],
        ["Security — Role guard", "Role check server-side via require_roles dependency; no client-supplied role trusted",  "✅ Pass"],
        ["Concurrency",           "SELECT FOR UPDATE inside DB transaction prevents race condition on last copy",          "✅ Pass"],
        ["Data integrity",        "available_copies never goes negative due to DB constraint + app-level check",          "✅ Pass"],
        ["Error handling",        "All rejection paths return descriptive detail fields with correct HTTP codes",         "✅ Pass"],
        ["Pydantic v2",           "ConfigDict(from_attributes=True) used correctly; no deprecated .from_orm() calls",    "✅ Pass"],
        ["Frontend",              "Borrow button hidden for guest/non-member; loading state prevents double submit",      "✅ Pass"],
        ["Test coverage",         "9 unit tests covering all ACs + error paths",                                         "✅ Pass"],
    ],
    col_widths=[2.0, 3.8, 1.2]
)

# --- STEP 7 ---
h2(doc, "Step 7 — Verify")
body(doc, "Agent: verify-test     Output: verify-test-result.md     Gate: ✅ PASSED (unit ≥90%, E2E ≥80%)", italic=True)

h3(doc, "Unit Tests (pytest)")
add_table(doc,
    ["Test Case", "Result"],
    [
        ["test_guest_cannot_borrow",            "✅ PASSED"],
        ["test_admin_cannot_borrow",            "✅ PASSED"],
        ["test_librarian_cannot_borrow",        "✅ PASSED"],
        ["test_successful_borrow",              "✅ PASSED"],
        ["test_cannot_borrow_unavailable_book", "✅ PASSED"],
        ["test_cannot_borrow_same_book_twice",  "✅ PASSED"],
        ["test_book_not_found",                 "✅ PASSED"],
        ["test_unauthenticated_cannot_borrow",  "✅ PASSED"],
        ["test_max_active_borrows_limit",       "✅ PASSED"],
    ],
    col_widths=[4.0, 1.0]
)
badge(doc, "Unit Tests", "9 / 9 passed — 100.0% (threshold ≥ 90%)", ok=True)

h3(doc, "E2E Tests (Playwright)")
add_table(doc,
    ["Suite", "Test Case", "Result"],
    [
        ["Deployment", "Backend API is accessible",      "✅ PASSED"],
        ["Deployment", "Frontend is accessible",         "✅ PASSED"],
        ["Deployment", "Can login with admin user",      "✅ PASSED"],
        ["Deployment", "Books page is accessible",       "✅ PASSED"],
        ["Deployment", "Borrow API endpoint exists",     "✅ PASSED"],
        ["Deployment", "Can access books via API",       "✅ PASSED"],
    ],
    col_widths=[1.5, 3.3, 1.2]
)
badge(doc, "E2E Tests", "6 / 6 passed — 100.0% (threshold ≥ 80%)", ok=True)
body(doc,
    "Note: The 'Borrow API endpoint exists' test was fixed during this session — "
    "the expected status set was updated to include 422 and 403, which the borrow endpoint "
    "correctly returns for a non-UUID book ID after TASK-10's UUID path param enforcement.",
    italic=True, size=9)

# --- STEP 8 ---
h2(doc, "Step 8 — PR Creation")
body(doc, "Agent: pr-creator     Output: PR #3 + CHANGELOG.md entry     Gate: ✅ Complete", italic=True)
add_table(doc,
    ["Item", "Detail"],
    [
        ["Branch",       "EPMCDMETST-40786"],
        ["PR Number",    "#3"],
        ["PR URL",       "https://github.com/JanpreetSingh/online-library-management/pull/3"],
        ["PR Title",     "feat(FR-3.1): Borrow a Book — member borrow endpoint with atomic availability management"],
        ["PR Sections",  "Summary · Changes Made · Test Evidence · Known Limitations · Reviewer Checklist"],
        ["CHANGELOG",    "## [Unreleased] — 2026-06-04 — Added FR-3.1 Borrow a Book (EPMCDMETST-40786)"],
        ["Status",       "Open — awaiting reviewer approval before merge + deployment"],
    ],
    col_widths=[1.5, 5.5]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 5 — FINAL ORCHESTRATOR OUTPUT
# ---------------------------------------------------------------------------

h1(doc, "5. Final Orchestrator Output")
body(doc,
    "The following table is the session summary produced by orchestrator-sdlc upon completion "
    "of all 8 steps for EPMCDMETST-40786.")

add_table(doc,
    ["REQ-ID", "Jira", "Title", "Refine Req", "Design", "Design Review", "Plan", "Impl", "Code Review", "Verify", "PR"],
    [
        ["FR-3.1", "EPMCDMETST-40786", "Borrow a Book", "✅", "✅", "✅", "✅", "✅", "✅", "15/15 ✅", "✅"],
    ],
    col_widths=[0.6, 1.7, 1.2, 0.75, 0.65, 1.0, 0.45, 0.45, 0.9, 0.75, 0.45]
)

body(doc, "\nNext step: once PR #3 is merged, run @orchestrator-deployment EPMCDMETST-40786 FR-3.1 to deploy and update Confluence documentation.", italic=True)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUT)
print(f"✅ Report saved → {OUT}")
